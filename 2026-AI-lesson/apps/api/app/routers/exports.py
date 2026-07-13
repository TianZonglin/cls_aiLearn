from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.datetime_utils import to_utc_iso
from app.core.db import get_db_session
from app.core.paths import ensure_runtime_dirs
from app.models.chat_sessions import ChatSession
from app.models.export_jobs import ExportJob
from app.repositories.chat import get_chat_session
from app.repositories.knowledge_bases import get_knowledge_base, utc_now
from app.schemas.exports import ExportJobResponse, ExportRequest

router = APIRouter(prefix="/export", tags=["export"])


def serialize_export_job(job: ExportJob) -> ExportJobResponse:
    return ExportJobResponse(
        id=job.id,
        format=job.format,
        status=job.status,
        output_path=job.output_path,
        download_url=f"/export/{job.id}/download" if job.output_path and job.status == "done" else None,
        error_message=job.error_message,
        created_at=to_utc_iso(job.created_at),
        finished_at=to_utc_iso(job.finished_at) if job.finished_at else None,
    )


def resolve_knowledge_base_names(payload: ExportRequest, session: Session) -> list[str]:
    if payload.knowledge_base_names:
        return payload.knowledge_base_names
    if not payload.knowledge_base_ids:
        return []
    return [
        kb.name
        for kb_id in payload.knowledge_base_ids
        for kb in [get_knowledge_base(session, kb_id)]
        if kb is not None
    ]


def build_export_lines(payload: ExportRequest, session: Session) -> list[str]:
    now = utc_now()
    knowledge_base_names = resolve_knowledge_base_names(payload, session)

    lines = [
        "# 知识库问答导出",
        "",
        f"- 导出时间：{to_utc_iso(now)}",
        f"- 所选知识库：{'、'.join(knowledge_base_names) if knowledge_base_names else '未提供'}",
    ]

    if payload.session_id:
        chat_session = get_chat_session(session, payload.session_id)
        if chat_session is not None:
            lines.append(f"- 会话 ID：{chat_session.id}")

    lines.extend(
        [
            "",
            "## 问题",
            "",
            payload.question,
            "",
            "## 回答",
            "",
            payload.answer,
            "",
            "## 来源列表",
            "",
        ]
    )

    if payload.citations:
        for index, citation in enumerate(payload.citations, start=1):
            lines.extend(
                [
                    f"{index}. 知识库：{citation.get('knowledge_base_name', '未知知识库')}",
                    f"   - 文档：{citation.get('document_name', '未知文档')}",
                    f"   - 位置：{citation.get('location_label', '未知位置')}",
                    f"   - 片段：{citation.get('snippet', '')}",
                ]
            )
    else:
        lines.append("当前没有可导出的来源。")

    lines.append("")
    return lines


def build_markdown_content(payload: ExportRequest, session: Session) -> str:
    return "\n".join(build_export_lines(payload, session))


def build_docx_file(path: Path, payload: ExportRequest, session: Session) -> None:
    knowledge_base_names = resolve_knowledge_base_names(payload, session)
    now = utc_now()

    doc = DocxDocument()
    doc.add_heading("知识库问答导出", level=1)
    doc.add_paragraph(f"导出时间：{to_utc_iso(now)}")
    doc.add_paragraph(f"所选知识库：{'、'.join(knowledge_base_names) if knowledge_base_names else '未提供'}")
    if payload.session_id:
        chat_session = get_chat_session(session, payload.session_id)
        if chat_session is not None:
            doc.add_paragraph(f"会话 ID：{chat_session.id}")

    doc.add_heading("问题", level=2)
    doc.add_paragraph(payload.question)
    doc.add_heading("回答", level=2)
    doc.add_paragraph(payload.answer)
    doc.add_heading("来源列表", level=2)

    if payload.citations:
        for index, citation in enumerate(payload.citations, start=1):
            doc.add_paragraph(
                f"{index}. 知识库：{citation.get('knowledge_base_name', '未知知识库')}",
                style="List Number",
            )
            doc.add_paragraph(f"文档：{citation.get('document_name', '未知文档')}")
            doc.add_paragraph(f"位置：{citation.get('location_label', '未知位置')}")
            doc.add_paragraph(f"片段：{citation.get('snippet', '')}")
    else:
        doc.add_paragraph("当前没有可导出的来源。")

    doc.save(path)


def create_export_job(payload: ExportRequest, export_format: str, session: Session) -> tuple[ExportJob, Path]:
    runtime_dirs = ensure_runtime_dirs()
    export_id = str(uuid4())
    now = utc_now()
    suffix = ".md" if export_format == "markdown" else ".docx"
    output_path = runtime_dirs["exports"] / f"{export_id}{suffix}"
    export_session_id = payload.session_id
    if export_session_id:
        if get_chat_session(session, export_session_id) is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Chat session not found.")
    else:
        export_session_id = session.execute(select(ChatSession.id).limit(1)).scalar_one_or_none()
        if export_session_id is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="导出需要关联一个已存在的会话，请先完成一次问答。",
            )

    job = ExportJob(
        id=export_id,
        session_id=export_session_id,
        format=export_format,
        status="processing",
        output_path=None,
        error_message=None,
        created_at=now,
        finished_at=None,
    )
    session.add(job)
    session.commit()
    session.refresh(job)
    return job, output_path


def finalize_export_job(job: ExportJob, output_path: Path, session: Session) -> ExportJobResponse:
    job.status = "done"
    job.output_path = str(output_path)
    job.finished_at = utc_now()
    session.add(job)
    session.commit()
    session.refresh(job)
    return serialize_export_job(job)


def fail_export_job(job: ExportJob, exc: Exception, session: Session) -> None:
    job.status = "failed"
    job.error_message = str(exc)
    job.finished_at = utc_now()
    session.add(job)
    session.commit()
    session.refresh(job)


@router.post("/markdown", response_model=ExportJobResponse, status_code=status.HTTP_201_CREATED)
def export_markdown(payload: ExportRequest, session: Session = Depends(get_db_session)) -> ExportJobResponse:
    job, output_path = create_export_job(payload, "markdown", session)
    try:
        output_path.write_text(build_markdown_content(payload, session), encoding="utf-8")
        return finalize_export_job(job, output_path, session)
    except Exception as exc:
        fail_export_job(job, exc, session)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc


@router.post("/docx", response_model=ExportJobResponse, status_code=status.HTTP_201_CREATED)
def export_docx(payload: ExportRequest, session: Session = Depends(get_db_session)) -> ExportJobResponse:
    job, output_path = create_export_job(payload, "docx", session)
    try:
        build_docx_file(output_path, payload, session)
        return finalize_export_job(job, output_path, session)
    except Exception as exc:
        fail_export_job(job, exc, session)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc


@router.get("/{export_id}", response_model=ExportJobResponse)
def get_export_job(export_id: str, session: Session = Depends(get_db_session)) -> ExportJobResponse:
    job = session.get(ExportJob, export_id)
    if job is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Export job not found.")
    return serialize_export_job(job)


@router.get("/{export_id}/download")
def download_export_file(export_id: str, session: Session = Depends(get_db_session)) -> FileResponse:
    job = session.get(ExportJob, export_id)
    if job is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Export job not found.")
    if job.status != "done" or not job.output_path:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Export file is not ready.")
    path = Path(job.output_path)
    if not path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Export file not found.")
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if job.format == "docx" else "text/markdown; charset=utf-8"
    return FileResponse(path=path, media_type=media_type, filename=path.name)
